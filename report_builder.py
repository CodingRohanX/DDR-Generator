import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMAGE_MARKER_RE = re.compile(r"\[Image:\s*([^\]]+)\]")


def _add_image_by_filename(doc, image_lookup, file_name):
    image_path = image_lookup.get(file_name)
    if not image_path or not os.path.exists(image_path):
        doc.add_paragraph("Image Not Available")
        return
    try:
        doc.add_picture(image_path, width=Inches(4.8))
        caption = doc.add_paragraph(file_name)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        doc.add_paragraph("Image Not Available")


def build_word_report(ddr_text, image_records, output_path):
    doc = Document()
    title = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared by: AI DDR System").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    image_lookup = {record["file_name"]: record["file_path"] for record in image_records}

    for raw_line in ddr_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=2)
            continue

        marker_match = IMAGE_MARKER_RE.search(line)
        if marker_match:
            before = IMAGE_MARKER_RE.sub("", line).strip("- ").strip()
            if before:
                doc.add_paragraph(before, style="List Bullet")
            _add_image_by_filename(doc, image_lookup, marker_match.group(1).strip())
            continue

        if line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Report saved to: {output_path}")